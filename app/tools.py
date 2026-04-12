import httpx
import asyncio
from datetime import datetime, timedelta
from app.config import COINGECKO_API_KEY

CHUNK_SIZE = 8000  # karakter per chunk
CHUNK_OVERLAP = 500  # overlap antar chunk biar konteks tidak putus

DEFILLAMA_BASE_URL = "https://api.llama.fi"


async def get_tvl_growth(protocol_name: str) -> dict:
    url = f"{DEFILLAMA_BASE_URL}/protocol/{protocol_name.lower()}"

    try:
        async with httpx.AsyncClient(timeout=10) as client:
            response = await client.get(url)

            if response.status_code == 404:
                return {"error": f"Protokol '{protocol_name}' tidak ditemukan di DefiLlama."}

            if response.status_code != 200:
                return {"error": f"DefiLlama API error: {response.status_code}"}

            data = response.json()

            tvl_data = data.get("tvl", [])

            if not tvl_data:
                return {"error": f"Data TVL untuk '{protocol_name}' tidak tersedia."}

            # Ambil TVL sekarang (index terakhir)
            current = tvl_data[-1]
            current_tvl = current["totalLiquidityUSD"]
            current_date = datetime.fromtimestamp(current["date"]).strftime("%d %b %Y")

            # Cari TVL 30 hari lalu
            target_date = datetime.now() - timedelta(days=30)
            past_entry = None

            for entry in reversed(tvl_data):
                entry_date = datetime.fromtimestamp(entry["date"])
                if entry_date <= target_date:
                    past_entry = entry
                    break

            if not past_entry:
                return {"error": f"Data 30 hari lalu untuk '{protocol_name}' tidak tersedia."}

            past_tvl = past_entry["totalLiquidityUSD"]
            past_date = datetime.fromtimestamp(past_entry["date"]).strftime("%d %b %Y")

            # Hitung growth
            if past_tvl == 0:
                return {"error": "TVL 30 hari lalu adalah 0, tidak bisa hitung growth."}

            growth = ((current_tvl - past_tvl) / past_tvl) * 100

            return {
                "protocol": data.get("name", protocol_name),
                "current_tvl": current_tvl,
                "current_date": current_date,
                "past_tvl": past_tvl,
                "past_date": past_date,
                "growth_percent": round(growth, 2)
            }

    except httpx.TimeoutException:
        return {"error": "DefiLlama API timeout. Coba lagi."}

    except Exception as e:
        return {"error": f"Error: {str(e)}"}


def format_tvl_result(result: dict) -> str:
    if "error" in result:
        return result["error"]

    growth = result["growth_percent"]
    arrow = "🟢 +" if growth >= 0 else "🔴 "

    current = result["current_tvl"]
    past = result["past_tvl"]

    def format_usd(value):
        if value >= 1_000_000_000:
            return f"${value / 1_000_000_000:.2f}B"
        elif value >= 1_000_000:
            return f"${value / 1_000_000:.2f}M"
        else:
            return f"${value:,.0f}"

    return (
        f"TVL {result['protocol']}:\n"
        f"- Sekarang ({result['current_date']}): {format_usd(current)}\n"
        f"- 30 hari lalu ({result['past_date']}): {format_usd(past)}\n"
        f"- Growth: {arrow}{growth}%"
    )

# ============================================
# CRYPTO MARKET TOOLS
# ============================================

COINGECKO_BASE_URL = "https://api.coingecko.com/api/v3"
COINGECKO_HEADERS = {"x-cg-demo-api-key": COINGECKO_API_KEY} if COINGECKO_API_KEY else {}

FEAR_GREED_URL = "https://api.alternative.me/fng"
DEXSCREENER_BASE_URL = "https://api.dexscreener.com"


def format_usd(value):
    """Format angka ke USD readable"""
    if not value or value == 0:
        return "\\\$0"
    if value >= 1_000_000_000_000:
        return f"${value / 1_000_000_000_000:.2f}T"
    elif value >= 1_000_000_000:
        return f"${value / 1_000_000_000:.2f}B"
    elif value >= 1_000_000:
        return f"${value / 1_000_000:.2f}M"
    elif value >= 1_000:
        return f"${value / 1_000:.2f}K"
    else:
        return f"${value:,.2f}"


def format_change(value):
    """Format persentase dengan emoji"""
    if value >= 0:
        return f"🟢 +{value}%"
    return f"🔴 {value}%"


async def get_global_market() -> dict:
    url = f"{COINGECKO_BASE_URL}/global"

    try:
        async with httpx.AsyncClient(timeout=10) as client:
            response = await client.get(url, headers=COINGECKO_HEADERS)

            if response.status_code != 200:
                return {"error": f"CoinGecko API error: {response.status_code}"}

            data = response.json().get("data", {})

            return {
                "total_market_cap_usd": data.get("total_market_cap", {}).get("usd", 0),
                "total_volume_24h_usd": data.get("total_volume", {}).get("usd", 0),
                "btc_dominance": round(data.get("market_cap_percentage", {}).get("btc", 0), 2),
                "eth_dominance": round(data.get("market_cap_percentage", {}).get("eth", 0), 2),
                "active_coins": data.get("active_cryptocurrencies", 0),
                "market_cap_change_24h": round(data.get("market_cap_change_percentage_24h_usd", 0), 2),
            }

    except httpx.TimeoutException:
        return {"error": "CoinGecko API timeout."}
    except Exception as e:
        return {"error": f"Error: {str(e)}"}


async def get_trending_coins() -> dict:
    url = f"{COINGECKO_BASE_URL}/search/trending"

    try:
        async with httpx.AsyncClient(timeout=10) as client:
            response = await client.get(url, headers=COINGECKO_HEADERS)

            if response.status_code != 200:
                return {"error": f"CoinGecko API error: {response.status_code}"}

            data = response.json()
            coins = data.get("coins", [])

            results = []
            for coin in coins[:10]:
                item = coin.get("item", {})
                data_block = item.get("data", {}) or {}
                price_change_block = data_block.get("price_change_percentage_24h", {}) or {}

                results.append({
                    "name": item.get("name", "?"),
                    "symbol": item.get("symbol", "?").upper(),
                    "market_cap_rank": item.get("market_cap_rank", "N/A"),
                    "score": item.get("score", 0),
                    "price_btc": item.get("price_btc", 0),
                    "price_change_24h": price_change_block.get("usd", 0),
                    "market_cap": data_block.get("market_cap", "N/A"),
                    "total_volume": data_block.get("total_volume", "N/A"),
                })

            return {"trending": results}

    except httpx.TimeoutException:
        return {"error": "CoinGecko API timeout."}
    except Exception as e:
        return {"error": f"Error: {str(e)}"}


async def get_top_movers() -> dict:
    url = f"{COINGECKO_BASE_URL}/coins/markets"
    params = {
        "vs_currency": "usd",
        "order": "market_cap_desc",
        "per_page": 100,
        "page": 1,
        "sparkline": "false",
        "price_change_percentage": "1h,24h,7d",
    }

    try:
        async with httpx.AsyncClient(timeout=15) as client:
            response = await client.get(url, headers=COINGECKO_HEADERS, params=params)

            if response.status_code != 200:
                return {"error": f"CoinGecko API error: {response.status_code}"}

            data = response.json()

            sorted_data = sorted(
                data,
                key=lambda x: x.get("price_change_percentage_24h") or 0,
                reverse=True
            )

            top_gainers = []
            for coin in sorted_data[:5]:
                top_gainers.append({
                    "name": coin.get("name", "?"),
                    "symbol": coin.get("symbol", "?").upper(),
                    "current_price": coin.get("current_price", 0),
                    "change_1h": round(coin.get("price_change_percentage_1h_in_currency") or 0, 2),
                    "change_24h": round(coin.get("price_change_percentage_24h") or 0, 2),
                    "change_7d": round(coin.get("price_change_percentage_7d_in_currency") or 0, 2),
                    "volume_24h": coin.get("total_volume", 0),
                    "market_cap": coin.get("market_cap", 0),
                    "market_cap_rank": coin.get("market_cap_rank", "N/A"),
                })

            top_losers = []
            for coin in sorted_data[-5:]:
                top_losers.append({
                    "name": coin.get("name", "?"),
                    "symbol": coin.get("symbol", "?").upper(),
                    "current_price": coin.get("current_price", 0),
                    "change_1h": round(coin.get("price_change_percentage_1h_in_currency") or 0, 2),
                    "change_24h": round(coin.get("price_change_percentage_24h") or 0, 2),
                    "change_7d": round(coin.get("price_change_percentage_7d_in_currency") or 0, 2),
                    "volume_24h": coin.get("total_volume", 0),
                    "market_cap": coin.get("market_cap", 0),
                    "market_cap_rank": coin.get("market_cap_rank", "N/A"),
                })

            return {"gainers": top_gainers, "losers": top_losers}

    except httpx.TimeoutException:
        return {"error": "CoinGecko API timeout."}
    except Exception as e:
        return {"error": f"Error: {str(e)}"}


async def get_coin_detail(coin_id: str) -> dict:
    url = f"{COINGECKO_BASE_URL}/coins/{coin_id.lower()}"
    params = {
        "localization": "false",
        "tickers": "false",
        "community_data": "false",
        "developer_data": "false",
        "sparkline": "false",
    }

    try:
        async with httpx.AsyncClient(timeout=10) as client:
            response = await client.get(url, headers=COINGECKO_HEADERS, params=params)

            if response.status_code == 404:
                return {"error": f"Coin '{coin_id}' tidak ditemukan."}
            if response.status_code != 200:
                return {"error": f"CoinGecko API error: {response.status_code}"}

            data = response.json()
            market = data.get("market_data", {})

            return {
                "name": data.get("name", "?"),
                "symbol": data.get("symbol", "?").upper(),
                "market_cap_rank": data.get("market_cap_rank", "N/A"),
                "current_price": market.get("current_price", {}).get("usd", 0),
                "market_cap": market.get("market_cap", {}).get("usd", 0),
                "total_volume": market.get("total_volume", {}).get("usd", 0),
                "high_24h": market.get("high_24h", {}).get("usd", 0),
                "low_24h": market.get("low_24h", {}).get("usd", 0),
                "change_24h": round(market.get("price_change_percentage_24h") or 0, 2),
                "change_7d": round(market.get("price_change_percentage_7d") or 0, 2),
                "change_30d": round(market.get("price_change_percentage_30d") or 0, 2),
                "ath": market.get("ath", {}).get("usd", 0),
                "ath_change": round(market.get("ath_change_percentage", {}).get("usd", 0), 2),
                "circulating_supply": market.get("circulating_supply", 0),
                "total_supply": market.get("total_supply", 0),
                "description": (data.get("description", {}).get("en", "") or "")[:500],
            }

    except httpx.TimeoutException:
        return {"error": "CoinGecko API timeout."}
    except Exception as e:
        return {"error": f"Error: {str(e)}"}


async def get_fear_greed() -> dict:
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            response = await client.get(f"{FEAR_GREED_URL}/?limit=7")

            if response.status_code != 200:
                return {"error": f"Fear & Greed API error: {response.status_code}"}

            data = response.json().get("data", [])

            if not data:
                return {"error": "Data Fear & Greed tidak tersedia."}

            current = data[0]
            yesterday = data[1] if len(data) > 1 else None
            week_ago = data[6] if len(data) > 6 else None

            result = {
                "value": int(current.get("value", 0)),
                "classification": current.get("value_classification", "?"),
                "timestamp": current.get("timestamp", ""),
            }

            if yesterday:
                result["yesterday_value"] = int(yesterday.get("value", 0))
                result["yesterday_class"] = yesterday.get("value_classification", "?")

            if week_ago:
                result["week_ago_value"] = int(week_ago.get("value", 0))
                result["week_ago_class"] = week_ago.get("value_classification", "?")

            return result

    except httpx.TimeoutException:
        return {"error": "Fear & Greed API timeout."}
    except Exception as e:
        return {"error": f"Error: {str(e)}"}


async def get_dex_trending() -> dict:
    url = f"{DEXSCREENER_BASE_URL}/token-boosts/latest"

    try:
        async with httpx.AsyncClient(timeout=10) as client:
            response = await client.get(url)

            if response.status_code != 200:
                return {"error": f"DexScreener API error: {response.status_code}"}

            data = response.json()

            if not data or not isinstance(data, list):
                return {"error": "Data DexScreener tidak tersedia."}

            results = []
            seen = set()

            for token in data[:20]:
                token_address = token.get("tokenAddress", "")
                if token_address in seen:
                    continue
                seen.add(token_address)

                results.append({
                    "chain": token.get("chainId", "?"),
                    "description": token.get("description", ""),
                    "url": token.get("url", ""),
                    "token_address": token_address,
                    "icon": token.get("icon", ""),
                })

                if len(results) >= 10:
                    break

            return {"dex_trending": results}

    except httpx.TimeoutException:
        return {"error": "DexScreener API timeout."}
    except Exception as e:
        return {"error": f"Error: {str(e)}"}


async def get_full_market_data() -> dict:
    results = await asyncio.gather(
        get_global_market(),
        get_top_movers(),
        get_trending_coins(),
        get_fear_greed(),
        get_dex_trending(),
        return_exceptions=True,
    )

    return {
        "global": results[0] if not isinstance(results[0], Exception) else {"error": str(results[0])},
        "movers": results[1] if not isinstance(results[1], Exception) else {"error": str(results[1])},
        "trending": results[2] if not isinstance(results[2], Exception) else {"error": str(results[2])},
        "fear_greed": results[3] if not isinstance(results[3], Exception) else {"error": str(results[3])},
        "dex": results[4] if not isinstance(results[4], Exception) else {"error": str(results[4])},
    }


def format_global_result(result: dict) -> str:
    if "error" in result:
        return result["error"]

    return (
        f"🌍 Global Crypto Market:\n\n"
        f"Market Cap: {format_usd(result.get('total_market_cap_usd', 0))}"
        f" ({format_change(result.get('market_cap_change_24h', 0))})\n"
        f"Volume 24h: {format_usd(result.get('total_volume_24h_usd', 0))}\n"
        f"BTC Dominance: {result.get('btc_dominance', 0)}%\n"
        f"ETH Dominance: {result.get('eth_dominance', 0)}%\n"
        f"Active Coins: {result.get('active_coins', 0):,}"
    )


def format_trending_result(result: dict) -> str:
    if "error" in result:
        return result["error"]

    coins = result.get("trending", [])
    if not coins:
        return "Tidak ada data trending."

    lines = ["🔥 Trending Coins (CoinGecko):\n"]

    for i, coin in enumerate(coins, 1):
        change = coin.get("price_change_24h", 0)
        change_str = format_change(round(change, 2)) if change else "N/A"
        rank = coin.get("market_cap_rank", "N/A")

        lines.append(
            f"{i}. ${coin['symbol']} ({coin['name']})"
            f"\n   Rank: #{rank} | 24h: {change_str}"
        )

    return "\n".join(lines)


def format_top_movers_result(result: dict) -> str:
    if "error" in result:
        return result["error"]

    lines = []

    gainers = result.get("gainers", [])
    if gainers:
        lines.append("📈 Top Gainers (24h):\n")
        for i, coin in enumerate(gainers, 1):
            lines.append(
                f"{i}. ${coin['symbol']} — {format_change(coin['change_24h'])}"
                f"\n   Price: ${coin['current_price']:,.4f} | Vol: {format_usd(coin['volume_24h'])}"
            )
        lines.append("")

    losers = result.get("losers", [])
    if losers:
        lines.append("📉 Top Losers (24h):\n")
        for i, coin in enumerate(losers, 1):
            lines.append(
                f"{i}. ${coin['symbol']} — {format_change(coin['change_24h'])}"
                f"\n   Price: ${coin['current_price']:,.4f} | Vol: {format_usd(coin['volume_24h'])}"
            )

    return "\n".join(lines)


def format_fear_greed_result(result: dict) -> str:
    if "error" in result:
        return result["error"]

    value = result["value"]

    if value <= 25:
        emoji = "😱"
    elif value <= 45:
        emoji = "😰"
    elif value <= 55:
        emoji = "😐"
    elif value <= 75:
        emoji = "😊"
    else:
        emoji = "🤑"

    text = f"{emoji} Fear & Greed Index: {value}/100 ({result['classification']})"

    if "yesterday_value" in result:
        text += f"\nKemarin: {result['yesterday_value']} ({result['yesterday_class']})"

    if "week_ago_value" in result:
        text += f"\n7 hari lalu: {result['week_ago_value']} ({result['week_ago_class']})"

    return text


def format_coin_detail_result(result: dict) -> str:
    if "error" in result:
        return result["error"]

    return (
        f"📊 {result['name']} (${result['symbol']}) — Rank #{result['market_cap_rank']}\n\n"
        f"Price: ${result['current_price']:,.4f}\n"
        f"24h: {format_change(result['change_24h'])} | "
        f"7d: {format_change(result['change_7d'])} | "
        f"30d: {format_change(result['change_30d'])}\n"
        f"24h Range: ${result['low_24h']:,.4f} — ${result['high_24h']:,.4f}\n"
        f"Market Cap: {format_usd(result['market_cap'])}\n"
        f"Volume 24h: {format_usd(result['total_volume'])}\n"
        f"ATH: ${result['ath']:,.4f} ({result['ath_change']}% from ATH)\n"
    )


def build_daily_pick_prompt(market_data: dict) -> str:
    """Format market data jadi teks data murni — tanpa instruksi"""
    parts = []

    # Global market
    g = market_data.get("global", {})
    if "error" not in g:
        parts.extend([
            "GLOBAL MARKET:",
            f"- Total Market Cap: {format_usd(g.get('total_market_cap_usd', 0))}",
            f"- 24h Market Cap Change: {g.get('market_cap_change_24h', 0)}%",
            f"- 24h Volume: {format_usd(g.get('total_volume_24h_usd', 0))}",
            f"- BTC Dominance: {g.get('btc_dominance', 0)}%",
            f"- ETH Dominance: {g.get('eth_dominance', 0)}%",
            ""
        ])

    # Fear & Greed
    fg = market_data.get("fear_greed", {})
    if "error" not in fg:
        fg_text = f"FEAR & GREED INDEX: {fg.get('value', '?')}/100 ({fg.get('classification', '?')})"
        if "yesterday_value" in fg:
            fg_text += f" | Yesterday: {fg['yesterday_value']} ({fg['yesterday_class']})"
        if "week_ago_value" in fg:
            fg_text += f" | 7d ago: {fg['week_ago_value']} ({fg['week_ago_class']})"
        parts.extend([fg_text, ""])

    # Top Gainers
    movers = market_data.get("movers", {})
    if "error" not in movers:
        gainers = movers.get("gainers", [])
        if gainers:
            parts.append("TOP GAINERS (24h):")
            for coin in gainers:
                vol_mcap_ratio = 0
                if coin.get('market_cap') and coin['market_cap'] > 0:
                    vol_mcap_ratio = round(coin['volume_24h'] / coin['market_cap'], 2)
                parts.append(
                    f"- ${coin['symbol']}: {coin['change_24h']}% (24h) | "
                    f"1h: {coin['change_1h']}% | 7d: {coin['change_7d']}% | "
                    f"Price: ${coin['current_price']} | "
                    f"Vol: {format_usd(coin['volume_24h'])} | "
                    f"MCap: {format_usd(coin['market_cap'])} | "
                    f"Vol/MCap: {vol_mcap_ratio}x | "
                    f"Rank: #{coin['market_cap_rank']}"
                )
            parts.append("")

        losers = movers.get("losers", [])
        if losers:
            parts.append("TOP LOSERS (24h):")
            for coin in losers:
                parts.append(
                    f"- ${coin['symbol']}: {coin['change_24h']}% (24h) | "
                    f"1h: {coin['change_1h']}% | 7d: {coin['change_7d']}% | "
                    f"Price: ${coin['current_price']} | "
                    f"Vol: {format_usd(coin['volume_24h'])} | "
                    f"MCap: {format_usd(coin['market_cap'])} | "
                    f"Rank: #{coin['market_cap_rank']}"
                )
            parts.append("")

    # Trending
    trending = market_data.get("trending", {})
    if "error" not in trending:
        coins = trending.get("trending", [])
        if coins:
            parts.append("TRENDING COINS (CoinGecko):")
            for coin in coins[:7]:
                change = coin.get('price_change_24h')
                change_str = f"{round(change, 2)}%" if change else "N/A"
                parts.append(
                    f"- ${coin['symbol']} ({coin['name']}): "
                    f"Rank #{coin['market_cap_rank']} | "
                    f"24h: {change_str}"
                )
            parts.append("")

    # DEX trending
    dex = market_data.get("dex", {})
    if "error" not in dex:
        dex_coins = dex.get("dex_trending", [])
        if dex_coins:
            parts.append("DEX TRENDING (DexScreener):")
            for token in dex_coins[:5]:
                chain = token.get('chain', '?')
                desc = token.get('description', '')[:80]
                parts.append(f"- [{chain}] {desc}")
            parts.append("")

    return "\n".join(parts)


# ============================================
# DOCUMENT READER TOOLS
# ============================================

import os
import csv
import json


def read_txt(file_path):
    """Baca file .txt"""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        with open(file_path, "r", encoding="latin-1") as f:
            return f.read()


def read_pdf(file_path):
    """Baca file .pdf"""
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(file_path)
        texts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text)

        return "\n".join(texts) if texts else "[PDF kosong atau tidak bisa dibaca]"
    except Exception as e:
        return f"[Error baca PDF: {e}]"


def read_docx(file_path):
    """Baca file .docx"""
    try:
        from docx import Document

        doc = Document(file_path)
        texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)

        # Baca juga tabel kalau ada
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    texts.append(row_text)

        return "\n".join(texts) if texts else "[DOCX kosong]"
    except Exception as e:
        return f"[Error baca DOCX: {e}]"


def read_csv_file(file_path):
    """Baca file .csv"""
    try:
        rows = []
        with open(file_path, "r", encoding="utf-8") as f:
            reader = csv.reader(f)
            for i, row in enumerate(reader):
                if i >= 100:  # Batasi 100 baris biar tidak boros token
                    rows.append(f"... (total baris dipotong di 100)")
                    break
                rows.append(" | ".join(row))

        return "\n".join(rows) if rows else "[CSV kosong]"
    except Exception as e:
        return f"[Error baca CSV: {e}]"


def read_json_file(file_path):
    """Baca file .json"""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        text = json.dumps(data, indent=2, ensure_ascii=False)

        # Batasi panjang
        if len(text) > 5000:
            text = text[:5000] + "\n... (dipotong, file terlalu panjang)"

        return text
    except Exception as e:
        return f"[Error baca JSON: {e}]"


def read_xlsx(file_path):
    """Baca file .xlsx"""
    try:
        from openpyxl import load_workbook

        wb = load_workbook(file_path, read_only=True)
        texts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            texts.append(f"--- Sheet: {sheet_name} ---")

            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= 100:  # Batasi 100 baris per sheet
                    texts.append("... (dipotong di 100 baris)")
                    break
                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip():
                    texts.append(row_text)

        wb.close()
        return "\n".join(texts) if texts else "[XLSX kosong]"
    except Exception as e:
        return f"[Error baca XLSX: {e}]"


# Mapping ekstensi ke fungsi
DOCUMENT_READERS = {
    ".txt": read_txt,
    ".pdf": read_pdf,
    ".docx": read_docx,
    ".csv": read_csv_file,
    ".json": read_json_file,
    ".xlsx": read_xlsx,
    ".xls": read_xlsx,
    ".log": read_txt,
    ".md": read_txt,
    ".py": read_txt,
    ".js": read_txt,
    ".html": read_txt,
    ".xml": read_txt,
    ".yaml": read_txt,
    ".yml": read_txt,
    ".env": read_txt,
    ".ini": read_txt,
    ".cfg": read_txt,
    ".sql": read_txt,
}

# Ekstensi yang didukung
SUPPORTED_EXTENSIONS = list(DOCUMENT_READERS.keys())


# ============================================
# DOCUMENT CHUNKING
# ============================================


def split_text_into_chunks(text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    """Pecah teks panjang jadi chunks dengan overlap"""
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size

        # Coba potong di akhir paragraf atau kalimat
        if end < len(text):
            # Cari newline terdekat dari posisi end
            newline_pos = text.rfind("\n", start + chunk_size - 1000, end)
            if newline_pos > start:
                end = newline_pos + 1
            else:
                # Cari titik terdekat
                dot_pos = text.rfind(". ", start + chunk_size - 1000, end)
                if dot_pos > start:
                    end = dot_pos + 2

        chunks.append(text[start:end].strip())

        # Mulai chunk berikutnya dengan overlap
        start = end - overlap
        if start < 0:
            start = 0

        # Safety: kalau start tidak maju, paksa maju
        if start <= (end - chunk_size):
            start = end

    return chunks


def extract_text_from_file(file_path):
    """Ekstrak teks dari file berdasarkan ekstensi — TANPA batasan karakter"""
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    if ext not in DOCUMENT_READERS:
        return None, f"Format {ext} belum didukung. Format yang didukung: {', '.join(SUPPORTED_EXTENSIONS)}"

    reader = DOCUMENT_READERS[ext]
    text = reader(file_path)

    if not text or not text.strip():
        return None, "Dokumen kosong atau tidak bisa dibaca."

    return text, None


# ============================================
# DOCUMENT CREATOR TOOLS
# ============================================

def safe_text_for_pdf(text):
    """Clean text agar aman untuk PDF built-in fonts"""
    replacements = {
        '\u2018': "'", '\u2019': "'",
        '\u201c': '"', '\u201d': '"',
        '\u2013': '-', '\u2014': '--',
        '\u2026': '...', '\u2022': '-',
        '\u00a0': ' ', '\u200b': '',
        '\u2003': ' ', '\u2002': ' ',
    }
    for k, v in replacements.items():
        text = text.replace(k, v)

    text = text.replace("**", "").replace("*", "")
    return text.encode('latin-1', 'replace').decode('latin-1')


def parse_title_from_content(content):
    """Ambil judul dari baris pertama yang diawali #"""
    lines = content.strip().split("\n")
    title = "Dokumen"
    body_start = 0

    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith("# ") and not stripped.startswith("## "):
            title = stripped[2:].strip().replace("**", "").replace("*", "")
            body_start = i + 1
            break

    body = "\n".join(lines[body_start:])
    return title, body


def create_pdf_file(content, file_path):
    """Buat file PDF dari konten terstruktur"""
    from fpdf import FPDF
    from datetime import datetime

    title, body = parse_title_from_content(content)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 18)
    pdf.multi_cell(0, 10, safe_text_for_pdf(title), align="C")
    pdf.ln(3)

    # Date
    pdf.set_font("Helvetica", "I", 9)
    pdf.cell(0, 6, f"Dibuat: {datetime.now().strftime('%d %B %Y, %H:%M')}", ln=True, align="C")
    pdf.ln(8)

    # Separator line
    pdf.set_draw_color(200, 200, 200)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(8)

    # Body
    for line in body.split("\n"):
        stripped = line.strip()
        safe_line = safe_text_for_pdf(stripped)

        if not stripped:
            pdf.ln(4)
        elif stripped.startswith("### "):
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 12)
            pdf.multi_cell(0, 7, safe_text_for_pdf(stripped[4:]))
            pdf.ln(2)
        elif stripped.startswith("## "):
            pdf.ln(6)
            pdf.set_font("Helvetica", "B", 14)
            pdf.multi_cell(0, 8, safe_text_for_pdf(stripped[3:]))
            pdf.ln(3)
        elif stripped.startswith("- "):
            pdf.set_font("Helvetica", "", 11)
            x = pdf.get_x()
            pdf.cell(8)
            pdf.multi_cell(0, 6, safe_text_for_pdf(f"- {stripped[2:]}"))
            pdf.ln(1)
        else:
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 6, safe_line)
            pdf.ln(2)

    pdf.output(file_path)
    return title


def create_docx_file(content, file_path):
    """Buat file DOCX dari konten terstruktur"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import datetime

    title, body = parse_title_from_content(content)

    doc = Document()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Dibuat: {datetime.now().strftime('%d %B %Y, %H:%M')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Body
    for line in body.split("\n"):
        stripped = line.strip()

        if not stripped:
            continue
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].replace("**", ""), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].replace("**", ""), level=2)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].replace("**", ""), style="List Bullet")
        else:
            clean = stripped.replace("**", "").replace("*", "")
            doc.add_paragraph(clean)

    doc.save(file_path)
    return title
